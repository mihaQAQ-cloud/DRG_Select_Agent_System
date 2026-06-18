"""
文档自动生成智能体（单文件实现）
功能：逐章调用 LLM 生成 SRS / 架构设计 / 测试文档，排版为 .docx，提交到文档系统。

修复说明（v1.1）：
  1. 消除"双文件"问题：原先 _save_and_submit 先把文档存为 data/outputs/<filename>
     （如 srs.docx），再 POST 给 doc_system，而 doc_system.submit 接口会再把上传
     的文件另存一份 data/outputs/<doc_id>_<filename>（如 2494d572_srs.docx）。
     这导致每次运行都产生两个文件：一个名称正常、一个带乱码前缀的"中间文件"。

     修复方案：_save_and_submit 改为先把文件存到系统临时目录，提交成功后
     doc_system 负责保管带前缀的正式副本；若提交失败，再把临时文件复制到
     data/outputs/ 兜底。运行后 data/outputs/ 里只会有 doc_system 生成的带
     doc_id 前缀的单一文件，不再出现重复。
"""
import os
import shutil
import tempfile
from datetime import datetime
from typing import Optional

import requests
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from agents.llm_service import LLMService

llm = LLMService()

DOC_SYSTEM_URL = "http://localhost:8001"

# SRS 章节定义：(章节名, LLM 生成提示词)
SRS_CHAPTERS = [
    (
        "1. 引言",
        "请为医保 DRG 入组智能体系统的 SRS 生成【引言】章节（约 300 字），"
        "包括：目的、项目范围、定义与缩略词（DRG/MDC/ADRG/CC/MCC/ICD-10/LLM/EMR）、参考资料。",
    ),
    (
        "2. 总体描述",
        "请为医保 DRG 入组智能体系统生成【总体描述】章节（约 300 字），"
        "包括：产品概述（四个智能体子系统）、"
        "用户特征（医疗编码员/医保管理人员/软件测试员/系统管理员）、假设与依赖。",
    ),
    (
        "3. 功能需求",
        "请为医保 DRG 入组智能体系统生成【功能需求】章节（约 500 字），分四个子节描述：\n"
        "3.1 DRG 入组智能体（病历解析、MDC 匹配、ADRG 分组、CC/MCC 评估、DRG 输出、入组说明、批量处理、异常提示）\n"
        "3.2 文档自动生成智能体（SRS/架构/测试文档生成、Word 格式输出、版本记录）\n"
        "3.3 测试用例生成智能体（正常/边界/异常场景、Excel 导出、覆盖度统计）\n"
        "3.4 虚拟文档系统（文档接收、目录管理、检索下载、REST API）。",
    ),
    (
        "4. 非功能需求",
        "请为医保 DRG 入组智能体系统生成【非功能需求】章节（约 200 字），包括：\n"
        "性能（单次入组 <= 2s，批量 100 条 <= 30s）、"
        "可靠性（LLM 不可用时自动降级为规则引擎）、"
        "安全性（API Key 本地存储）、"
        "可维护性（规则文件热更新，无需重启服务）。",
    ),
]

# 架构设计章节
ARCH_CHAPTERS = [
    (
        "1. 系统总体架构",
        "请为医保 DRG 入组智能体系统生成【系统总体架构】章节（约 300 字），"
        "描述多智能体协作架构：main.py 主控编排器、四个智能体、共享 LLM 服务、SQLite 存储。",
    ),
    (
        "2. 模块划分",
        "请生成【模块划分】章节（约 300 字），说明各文件职责：\n"
        "main.py / config.yaml / requirements.txt / agents/drg_agent.py / "
        "agents/doc_agent.py / agents/test_agent.py / agents/llm_service.py / "
        "doc_system.py / rules/drg_rules.json。",
    ),
    (
        "3. 接口设计",
        "请生成【接口设计】章节（约 300 字），列出所有 REST 接口：\n"
        "DRG 入组智能体（端口 8000）：POST /api/drg/classify, POST /api/drg/batch, "
        "GET /api/drg/reload, GET /api/health\n"
        "虚拟文档系统（端口 8001）：POST /api/docs/submit, GET /api/docs/list, "
        "GET /api/docs/{doc_id}/download, GET /api/docs/search。",
    ),
    (
        "4. 数据库设计",
        "请生成【数据库设计】章节（约 200 字），描述 SQLite 数据库两张表：\n"
        "emr_records（病历记录）和 drg_results（分组结果）的字段设计，以及 documents 表的设计。",
    ),
]

# 测试文档章节
TEST_CHAPTERS = [
    (
        "1. 测试策略",
        "请生成【测试策略】章节（约 200 字），说明三种测试场景：\n"
        "正常场景（程序化生成标准入组病历）、"
        "边界场景（CC/MCC 有无、排除表边界）、"
        "异常场景（ICD 编码错误、字段缺失、格式异常）。",
    ),
    (
        "2. 测试用例模板",
        "请生成【测试用例模板】章节（约 200 字），说明 Excel 测试用例格式：\n"
        "列定义（用例ID/类型/输入病历JSON/期望DRG组号/说明）及填写示例。",
    ),
    (
        "3. 测试覆盖度要求",
        "请生成【测试覆盖度】章节（约 150 字），说明覆盖度统计方法：\n"
        "统计测试用例对每个 ADRG x 严重程度（MCC/CC/无）组合的覆盖情况，标注未覆盖项。",
    ),
]


def _add_title_page(doc: Document, title: str, version: str) -> None:
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.runs[0]
    run.font.size = Pt(22)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "版本：" + version + "    生成时间：" + datetime.now().strftime("%Y-%m-%d %H:%M:%S") + "\n"
        "华东理工大学 自然语言处理与大数据挖掘实验室\n软件工程课程大作业"
    ).font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    doc.add_page_break()


def _generate_doc(chapters: list, title: str, context: str) -> Document:
    doc = Document()
    _add_title_page(doc, title, "1.0")
    for chapter_name, prompt in chapters:
        full_prompt = prompt + "\n\n系统背景参考（节选）：" + context[:400]
        content = llm.chat(full_prompt)
        doc.add_heading(chapter_name, level=1)
        for para in content.split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
    return doc


def _save_and_submit(
    doc: Document,
    filename: str,
    doc_type: str,
    title: str,
    version: str,
) -> str:
    """
    将文档先写到临时文件，然后 POST 给虚拟文档系统。
    doc_system 负责将其存为 data/outputs/<doc_id>_<filename>，
    本函数不再额外写一份到 data/outputs/，从而消除双文件问题。
    若提交失败，才将临时文件复制到 data/outputs/<filename> 兜底。
    """
    os.makedirs("data/outputs", exist_ok=True)

    # 写入临时文件（与最终输出目录无关）
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
    doc.save(tmp_path)

    try:
        with open(tmp_path, "rb") as f:
            resp = requests.post(
                DOC_SYSTEM_URL + "/api/docs/submit",
                files={
                    "file": (
                        filename,
                        f,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
                data={"doc_type": doc_type, "title": title, "version": version},
                timeout=10,
            )
        doc_id = resp.json().get("doc_id", "unknown")
        print("[DocAgent] " + doc_type + " 文档已提交，doc_id=" + doc_id)
        return doc_id
    except Exception as exc:
        # 提交失败时兜底：把临时文件复制到 data/outputs/<filename>
        fallback_path = os.path.join("data/outputs", filename)
        shutil.copy2(tmp_path, fallback_path)
        print(
            "[DocAgent] 文档系统提交失败（" + str(exc) + "），"
            "文档已保存至 " + fallback_path
        )
        return "local-saved"
    finally:
        # 无论成功与否都删除临时文件
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


def generate(requirement_text: str) -> str:
    """
    根据需求描述文本生成 SRS 文档，提交到文档系统并返回 doc_id。对应 FA-02-01。
    """
    doc = _generate_doc(
        SRS_CHAPTERS,
        "医保 DRG 入组智能体系统  软件需求规格说明书（SRS）",
        requirement_text,
    )
    return _save_and_submit(doc, "srs.docx", "SRS", "软件需求规格说明书", "1.0")


def generate_arch(context: str = "") -> str:
    """生成架构设计文档，对应 FA-02-02。"""
    doc = _generate_doc(
        ARCH_CHAPTERS,
        "医保 DRG 入组智能体系统  系统架构设计文档",
        context,
    )
    return _save_and_submit(doc, "architecture.docx", "Architecture", "系统架构设计文档", "1.0")


def generate_test_doc(context: str = "") -> str:
    """生成测试文档框架，对应 FA-02-03。"""
    doc = _generate_doc(
        TEST_CHAPTERS,
        "医保 DRG 入组智能体系统  测试文档",
        context,
    )
    return _save_and_submit(doc, "test_doc.docx", "TestDoc", "测试文档", "1.0")
